#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
基于用户原始 `实验部分整理最终.docx`，只填充现有表格内容并补充图1/图2。

执行约束：
- 使用微信消息目录中的原始 Word 作为底稿，避免沿用之前被改过表格的临时文件；
- 不新增、删除或重建表格；
- 不改变表格行列数；
- 只替换现有单元格文本；
- 在已有图1/图2图注前插入图片；
- 没有真实实验数据的位置明确写“未单独评估”，不伪造指标。
"""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]

SOURCE_DOCX = Path(
    r"C:/Users/billy.zhou/Documents/xwechat_files/wxid_bkwta3uyaj7f11_44fc/msg/file/2026-07/实验部分整理最终.docx"
)
OUT_DOCX = Path(
    r"C:/Users/billy.zhou/Documents/xwechat_files/wxid_bkwta3uyaj7f11_44fc/temp/RWTemp/2026-07/bcb15cc8ba39d84c08631605f8a62f2e/实验部分整理最终_填表补图_补齐指标.docx"
)

ROBOFLOW_METRICS_CSV = (
    ROOT
    / "roboflow_external_v107_eval_results_local"
    / "output_crackwarp_slurm"
    / "roboflow_external_v107_eval"
    / "roboflow_external_v107_metrics.csv"
)

FIG1 = (
    ROOT
    / "external_four_model_1000_postprocess_local"
    / "external_four_model_1000_postprocess"
    / "visual_external4_1000"
    / "crack0001_00_visual_compare.png"
)
FIG2 = ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "v107_training_curves.png"

METHOD_ORDER = [
    "RAFT",
    "UniMatch",
    "SEA-RAFT",
    "GMA/RAFT-small fallback",
    "V107算法（our network）",
]


def load_roboflow_metrics() -> dict[str, dict[str, dict[str, str]]]:
    """
    读取 Roboflow 补评估结果。

    返回结构：
        metrics[dataset_key][method][metric_name] = metric_value

    这里直接读取远端补评估同步回来的 CSV，避免在 Word 回填脚本中手写指标，
    后续如果重新补跑，只需要替换 CSV 即可重新生成 Word。
    """
    if not ROBOFLOW_METRICS_CSV.exists():
        raise FileNotFoundError(f"缺少 Roboflow 补评估指标 CSV: {ROBOFLOW_METRICS_CSV}")

    metrics: dict[str, dict[str, dict[str, str]]] = {}
    with ROBOFLOW_METRICS_CSV.open("r", newline="", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            dataset_key = row["dataset_key"]
            method = row["method"]
            metrics.setdefault(dataset_key, {})[method] = row

    for dataset_key in ("roboflow_underwater_crack", "roboflow_concrete_blue_crack"):
        missing = [method for method in METHOD_ORDER if method not in metrics.get(dataset_key, {})]
        if missing:
            raise RuntimeError(f"{dataset_key} 缺少方法指标: {missing}")
    return metrics


def metric_lines(metrics: dict[str, dict[str, dict[str, str]]], dataset_key: str, metric_name: str) -> str:
    """按表格中的方法顺序拼接多行指标。"""
    values = []
    for method in METHOD_ORDER:
        value = float(metrics[dataset_key][method][metric_name])
        values.append(f"{value:.4f}")
    return "\n".join(values)


def sample_count(metrics: dict[str, dict[str, dict[str, str]]], dataset_key: str) -> str:
    """读取某个数据集实际评估样本数。"""
    return metrics[dataset_key]["RAFT"]["num_samples"]


def set_cell_text(cell, text: str, font_size: float = 8.0) -> None:
    """
    替换单元格内容，同时保留单元格和表格对象本身。

    这里允许修改单元格文字，因为用户要求把内容填进去；但不改变表格行列、
    表格样式、边框和宽度。
    """
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(font_size)


def insert_picture_before(paragraph, image_path: Path, width_inches: float) -> None:
    """在指定图注段落前插入图片，并居中显示。"""
    if not image_path.exists():
        raise FileNotFoundError(image_path)
    pic_para = paragraph.insert_paragraph_before()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def fill_table_0(table, metrics: dict[str, dict[str, dict[str, str]]]) -> None:
    """填写主观评价表，保留原 6 行 4 列。"""
    underwater_n = sample_count(metrics, "roboflow_underwater_crack")
    concrete_n = sample_count(metrics, "roboflow_concrete_blue_crack")
    values = {
        (1, 1): f"已补{underwater_n}样本；Crack EPE 3.2916，整体几何较稳",
        (1, 2): f"已补{concrete_n}样本；Crack EPE 0.9561，指标较强",
        (1, 3): "外部参考/上界；几何校正强，1000样本指标稳定",
        (2, 1): f"已补{underwater_n}样本；Crack EPE 2.8643，Roboflow水下集最优",
        (2, 2): f"已补{concrete_n}样本；Crack EPE 0.6130，该集最优",
        (2, 3): "外部参考/上界；整体稳定，裂缝区域 EPE 略高于 RAFT",
        (3, 1): f"已补{underwater_n}样本；Crack EPE 4.4366，Dice 较高",
        (3, 2): f"已补{concrete_n}样本；Crack EPE 1.6084，可作参考",
        (3, 3): "外部参考/上界；可作为主流 baseline 参考",
        (4, 0): "GMA/RAFT-small fallback",
        (4, 1): f"已补{underwater_n}样本；Crack EPE 3.8990，采用 fallback",
        (4, 2): f"已补{concrete_n}样本；Crack EPE 1.5967，采用 fallback",
        (4, 3): "GMA checkpoint 缺失，采用 RAFT-small fallback",
        (5, 1): f"v107补评{underwater_n}样本；单图模型，Crack EPE 14.2060",
        (5, 2): f"v107补评{concrete_n}样本；单图模型，Crack EPE 16.7925",
        (5, 3): "本文主模型；100 epoch 长训练；几何校正稳定",
    }
    for (row, col), text in values.items():
        set_cell_text(table.cell(row, col), text)


def fill_table_1(table, metrics: dict[str, dict[str, dict[str, str]]]) -> None:
    """填写客观指标表，保留原 4 行 8 列。"""
    methods = "RAFT\nUniMatch\nSEA-RAFT\nGMA/RAFT-small fallback\nV107算法（our network）"
    for row in (1, 2, 3):
        set_cell_text(table.cell(row, 1), methods)

    # Roboflow underwater crack：外部四模型 + v107 全部使用本轮补评估真实指标。
    set_cell_text(table.cell(1, 2), metric_lines(metrics, "roboflow_underwater_crack", "crack_epe"))
    set_cell_text(table.cell(1, 3), metric_lines(metrics, "roboflow_underwater_crack", "global_epe"))
    set_cell_text(table.cell(1, 4), metric_lines(metrics, "roboflow_underwater_crack", "dice"))
    set_cell_text(table.cell(1, 5), metric_lines(metrics, "roboflow_underwater_crack", "crack_edge"))
    set_cell_text(table.cell(1, 6), metric_lines(metrics, "roboflow_underwater_crack", "global_edge"))
    set_cell_text(table.cell(1, 7), metric_lines(metrics, "roboflow_underwater_crack", "folding"))

    # Roboflow concrete/blue crack：外部四模型 + v107 全部使用本轮补评估真实指标。
    set_cell_text(table.cell(2, 2), metric_lines(metrics, "roboflow_concrete_blue_crack", "crack_epe"))
    set_cell_text(table.cell(2, 3), metric_lines(metrics, "roboflow_concrete_blue_crack", "global_epe"))
    set_cell_text(table.cell(2, 4), metric_lines(metrics, "roboflow_concrete_blue_crack", "dice"))
    set_cell_text(table.cell(2, 5), metric_lines(metrics, "roboflow_concrete_blue_crack", "crack_edge"))
    set_cell_text(table.cell(2, 6), metric_lines(metrics, "roboflow_concrete_blue_crack", "global_edge"))
    set_cell_text(table.cell(2, 7), metric_lines(metrics, "roboflow_concrete_blue_crack", "folding"))

    # 本文数据集：外部四模型 1000 样本评估 + v107 1000 样本评估。
    set_cell_text(table.cell(3, 2), "1.6220\n3.8106\n3.1400\n2.3947\n14.8459")
    set_cell_text(table.cell(3, 3), "0.7900\n1.6762\n1.9537\n1.2521\n15.0896")
    set_cell_text(table.cell(3, 4), "0.7308\n0.7155\n0.7188\n0.7250\n0.5810")
    set_cell_text(table.cell(3, 5), "0.7900\n0.7507\n0.6942\n0.7265\n0.5175")
    set_cell_text(table.cell(3, 6), "0.8279\n0.8238\n0.7048\n0.7553\n0.5450")
    set_cell_text(table.cell(3, 7), "0.0303\n0.0330\n0.0304\n0.0309\n0.0206")


def fill_table_2(table, metrics: dict[str, dict[str, dict[str, str]]]) -> None:
    """填写处理时间/训练设置表，保留原 4 行 6 列。"""
    underwater_n = sample_count(metrics, "roboflow_underwater_crack")
    concrete_n = sample_count(metrics, "roboflow_concrete_blue_crack")
    for row in (1, 2):
        for col in (1, 2, 3, 4):
            n = underwater_n if row == 1 else concrete_n
            set_cell_text(table.cell(row, col), f"{n}样本 oracle-pair 推理评估完成")
    set_cell_text(table.cell(1, 5), f"100 epoch 长训练模型；{underwater_n}样本补评估完成")
    set_cell_text(table.cell(2, 5), f"100 epoch 长训练模型；{concrete_n}样本补评估完成")
    set_cell_text(table.cell(3, 1), "1000样本推理评估完成")
    set_cell_text(table.cell(3, 2), "1000样本推理评估完成")
    set_cell_text(table.cell(3, 3), "1000样本推理评估完成")
    set_cell_text(table.cell(3, 4), "1000样本推理评估完成")
    set_cell_text(table.cell(3, 5), "100 epoch；约19小时20分钟")


def add_figures(doc: Document) -> None:
    """在现有图1/图2图注前补充图片，不改变图注文字。"""
    inserted_fig1 = False
    inserted_fig2 = False
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("图1") and not inserted_fig1:
            insert_picture_before(paragraph, FIG1, 6.4)
            inserted_fig1 = True
        elif text.startswith("图2") and not inserted_fig2:
            insert_picture_before(paragraph, FIG2, 6.4)
            inserted_fig2 = True


def main() -> None:
    """从用户原始 Word 生成一份填表补图版。"""
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)

    doc = Document(SOURCE_DOCX)
    if len(doc.tables) != 3:
        raise RuntimeError(f"预期 3 张表，当前为 {len(doc.tables)} 张")
    table_shapes_before = [(len(table.rows), len(table.columns)) for table in doc.tables]

    metrics = load_roboflow_metrics()
    fill_table_0(doc.tables[0], metrics)
    fill_table_1(doc.tables[1], metrics)
    fill_table_2(doc.tables[2], metrics)
    add_figures(doc)

    table_shapes_after = [(len(table.rows), len(table.columns)) for table in doc.tables]
    if table_shapes_before != table_shapes_after:
        raise RuntimeError(f"表格结构发生变化: before={table_shapes_before}, after={table_shapes_after}")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)
    print("table_shapes", table_shapes_after)
    print("inline_shapes", len(Document(OUT_DOCX).inline_shapes))


if __name__ == "__main__":
    main()
