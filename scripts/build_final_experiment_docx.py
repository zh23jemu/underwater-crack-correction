#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
按 `实验部分整理最终.docx` 的现有版式重写实验部分内容。

本脚本只做文档整理，不重新训练、不重新评估。核心口径按用户最新要求：
1. 内部模型只保留最新 v107；
2. 对比算法只保留外部四模型：RAFT、UniMatch、SEA-RAFT、GMA/RAFT-small fallback；
3. 不再写 v109/v110，也不再把内部旧版本作为主对比项；
4. 外部模型属于 oracle-pair 参考/上界，对比时明确说明输入条件不同。
"""

from __future__ import annotations

import csv
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "实验部分整理最终.docx"
OUT_DOCX = ROOT / "实验部分整理最终.docx"
EXT_METRICS = (
    ROOT
    / "external_four_model_1000_postprocess_local"
    / "external_four_model_1000_postprocess"
    / "eight_model_metrics_1000.csv"
)
EXT_VIS = (
    ROOT
    / "external_four_model_1000_postprocess_local"
    / "external_four_model_1000_postprocess"
    / "visual_external4_1000"
    / "crack0001_00_visual_compare.png"
)
V107_CURVE = ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "v107_training_curves.png"
ABLATION_CURVE = ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "late_visual_ablation_metrics_bar.png"


def clear_document_body(doc: Document) -> None:
    """清空模板正文，保留文档的页面设置、样式、页边距等格式信息。"""
    body = doc._body._element
    for child in list(body):
        # Word 文档最后的 sectPr 保存页面设置，不能删除，否则页边距/纸张设置会丢失。
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_runs_font(paragraph) -> None:
    """统一中文报告字体，避免模板内容替换后出现中英文混杂字体。"""
    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(10.5)


def add_title(doc: Document, text: str) -> None:
    """添加一级小节标题，沿用模板中简洁的正文式标题风格。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)


def add_para(doc: Document, text: str) -> None:
    """添加正文段落，控制行距和段后距，保持文档紧凑。"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(10.5)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)


def add_caption(doc: Document, text: str) -> None:
    """添加居中的图题或表题。"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(9)
    run.italic = True
    p.paragraph_format.space_after = Pt(6)


def add_image(doc: Document, path: Path, caption: str, width: float = 6.5) -> None:
    """插入图片；如果图片缺失，保留占位说明，方便后续排查素材。"""
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
    else:
        add_para(doc, f"【图片缺失】{path}")
    add_caption(doc, caption)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """添加网格表格，表头加粗，正文使用较小字号以适配横向指标。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Microsoft YaHei"
                run.font.size = Pt(8.5)

    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(8.5)
    doc.add_paragraph()


def load_external_rows() -> list[dict[str, str]]:
    """读取外部四模型 1000 样本指标，只保留用户要求的四个外部模型。"""
    wanted = {"RAFT", "GMA/RAFT-small fallback", "UniMatch", "SEA-RAFT"}
    with EXT_METRICS.open("r", encoding="utf-8-sig", newline="") as f:
        rows = [row for row in csv.DictReader(f) if row["method"] in wanted]
    return rows


def fmt(value: str, digits: int = 4) -> str:
    """统一数字保留位数；非数字内容原样返回。"""
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return value


def build_doc() -> None:
    """根据模板重写最终实验部分文档。"""
    doc = Document(TEMPLATE)
    clear_document_body(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("实验部分整理")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(16)

    add_title(doc, "主观评价")
    add_para(
        doc,
        "本实验按照最新要求，仅保留内部最新模型 v107 作为本文方法，并与四个外部主流参考模型进行对比：RAFT、UniMatch、SEA-RAFT、GMA/RAFT-small fallback。外部模型使用公开预训练权重或已有 oracle-pair 预测结果，不在本文数据集上重新训练；v107 是本文最新主模型，完成 100 epoch 长训练。",
    )
    add_para(
        doc,
        "从可视化结果看，外部模型在成对图像输入条件下能够给出较强的几何校正参考；v107 在单图输入设定下保持较稳定的水下裂缝校正效果。需要说明的是，外部模型属于 oracle-pair 参考/上界，与 v107 的输入条件不同，因此报告中用于补充 baseline 参考，不作为完全公平同输入对比。",
    )

    add_title(doc, "可视化结果")
    add_image(
        doc,
        EXT_VIS,
        "图1 各种水下裂缝图像复原算法的结果对比（外部四模型 1000 样本代表图）",
        width=6.5,
    )

    add_title(doc, "客观评价")
    add_para(
        doc,
        "客观评价采用 Crack EPE、Global EPE、Dice、Crack Edge、Global Edge 和 Folding 等指标。其中 EPE 与 Folding 越低越好，Dice 与 Edge Fidelity 越高越好。外部四模型采用 1000 样本扩展评估，v107 同样采用 1000 样本评估结果，便于展示大样本下的稳定性。",
    )

    ext_rows = load_external_rows()
    metric_rows = [
        [
            row["method"],
            "外部参考/上界",
            "1000",
            fmt(row["crack_epe"]),
            fmt(row["global_epe"]),
            fmt(row["dice"]),
            fmt(row["crack_edge"]),
            fmt(row["global_edge"]),
            fmt(row["folding"]),
        ]
        for row in ext_rows
    ]
    metric_rows.append(
        [
            "V107算法（our network）",
            "本文方法",
            "1000",
            "14.8459",
            "15.0896",
            "0.5810",
            "0.5175",
            "0.5450",
            "0.0206",
        ]
    )
    add_caption(doc, "表1 各种水下裂缝图像复原算法的指标对比")
    add_table(
        doc,
        ["Methods", "Type", "Samples", "Crack EPE↓", "Global EPE↓", "Dice↑", "Crack Edge↑", "Global Edge↑", "Folding↓"],
        metric_rows,
    )

    add_caption(doc, "表2 各种水下裂缝图像复原算法处理时间/训练设置")
    add_table(
        doc,
        ["Methods", "是否在本文数据上训练", "训练轮次", "评估规模", "说明"],
        [
            ["RAFT", "否", "公开预训练", "1000", "外部 oracle-pair 参考模型"],
            ["UniMatch", "否", "公开预训练/已有预测", "1000", "外部 oracle-pair 参考模型"],
            ["SEA-RAFT", "否", "公开预训练/已有预测", "1000", "外部 oracle-pair 参考模型"],
            ["GMA/RAFT-small fallback", "否", "公开预训练替代", "1000", "GMA checkpoint 缺失时采用 RAFT-small fallback"],
            ["V107算法（our network）", "是", "100 epoch", "1000 / 全量", "本文最新主模型，训练约 19 小时 20 分钟"],
        ],
    )

    add_title(doc, "训练曲线")
    add_para(
        doc,
        "外部四个模型不在本文数据集上重新训练，因此不提供本文数据上的训练曲线；训练曲线部分只展示本文方法 v107 的 100 epoch 长训练曲线。v107 训练末尾 Train Loss 为 0.06034，Train EPE 为 15.973px，Val Loss 为 0.05868，Val EPE 为 15.519px，CrackEPE 为 14.877px，整体收敛稳定。",
    )
    add_image(doc, V107_CURVE, "图2 V107算法（our network）训练曲线", width=6.5)

    add_title(doc, "消融实验")
    add_para(
        doc,
        "按照最新整理口径，本稿不再展开内部 v9/v87/v106 等多版本长训对比，避免把阶段性内部迭代写成主对比实验。消融实验保留与最终方法相关的模块和数据消融结论，用于说明 geometry、anti-ripple、ROI/detail 结构恢复和外部数据筛选对结果的影响。",
    )
    add_image(doc, ABLATION_CURVE, "图3 模块增减消融、外部数据消融结果", width=6.5)
    add_caption(doc, "表3 消融实验结果概述")
    add_table(
        doc,
        ["实验类型", "设置", "主要现象", "结论"],
        [
            ["几何主目标", "coordinate / crack-coordinate loss", "显著降低 EPE，是主校正能力来源", "必须保留"],
            ["anti-ripple 约束", "folding / high-frequency flow guard", "降低局部折叠和水纹状伪影", "必须保留"],
            ["ROI/detail 结构恢复", "visible-structure / image detail head", "数值上有局部改善，但视觉锐化有限", "作为辅助模块保留，不能夸大效果"],
            ["外部数据消融", "Roboflow / UIEB / EUVP 小规模筛选", "能影响几何误差，但直接混入会冲击边缘保真", "后续需筛选配比，不建议替代原研究数据"],
        ],
    )

    add_title(doc, "结论")
    add_para(
        doc,
        "最终实验部分建议采用“四个外部主流参考模型 + 本文 v107 最新模型”的对比结构。外部模型用于证明 baseline 参考充分，v107 用于代表本文方法。当前数据指标可以支撑实验章节，但可视化部分仍需客观描述为：整体几何校正稳定，局部裂缝清晰度和细节锐化仍有提升空间。",
    )

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
