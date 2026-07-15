#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成面向客户的“实验部分”Word 文档。

文档结构严格对应客户截图中的 4 点要求：
1. 各种水下裂缝图像复原算法的可视化结果对比图；
2. 各种算法处理后图像质量指标的比较；
3. 算法复原水下裂缝图像训练曲线、时间；
4. 消融实验，模块增减消融、外部数据消融结果。

脚本只读取已经生成/同步到本地的实验结果、曲线和代表图，不重新训练模型。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "水下裂缝复原算法实验部分整理.docx"


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    """添加统一格式的标题。"""
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT


def add_paragraph(doc: Document, text: str) -> None:
    """添加正文段落，并统一中文报告常用行距和字号。"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(10.5)
    para.paragraph_format.line_spacing = 1.25
    para.paragraph_format.space_after = Pt(6)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """添加带表头的网格表格。"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(9)

    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            cell = row.cells[idx]
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(9)

    doc.add_paragraph()


def add_image(doc: Document, path: Path, caption: str, width: float = 6.6) -> None:
    """
    插入图片和图注。

    如果个别素材不存在，文档中会保留文字说明，避免生成流程中断。
    """
    if not path.exists():
        add_paragraph(doc, f"【图片缺失】{caption}：{path}")
        return
    doc.add_picture(str(path), width=Inches(width))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(9)
        run.italic = True


def setup_document() -> Document:
    """初始化 Word 页面、边距和默认字体。"""
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.72)
    section.right_margin = Inches(0.72)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10.5)
    return doc


def add_visual_section(doc: Document) -> None:
    """第 1 部分：可视化结果对比。"""
    add_heading(doc, "1. 各种水下裂缝图像复原算法的可视化结果对比图", 1)
    add_paragraph(
        doc,
        "本节展示不同阶段算法在代表样本上的复原效果。对比图中包含输入图像、GT/参考校正结果、传统或外部参考方法、早期主线模型、不同消融版本以及后期局部结构恢复结果。整体看，v107 作为最新主模型在几何误差和稳定性上显著优于原始基线；但局部裂缝纹理锐度仍是后续可优化点。",
    )
    add_image(
        doc,
        ROOT
        / "experiment_submission_package_2026-07-09_user_ready_v2"
        / "02_comparison_experiments"
        / "selected_compare_visuals"
        / "crack0030_00_visual_compare.png",
        "图1 代表样本的多算法复原效果对比（输入、GT、外部参考方法及内部消融版本）",
        width=6.6,
    )
    add_image(
        doc,
        ROOT / "client_visualization_report_2026-07-09" / "images" / "v106_crack0052_03_roi01_detail_panel.png",
        "图2 后期 detail-head / ROI 局部细节复核结果",
        width=6.6,
    )
    add_image(
        doc,
        ROOT / "comparison_user_original_vs_v107_assets" / "metric_comparison.png",
        "图3 用户原始 best_epe_v2 与 v107 最新主模型的核心指标对比",
        width=6.3,
    )
    add_image(
        doc,
        ROOT
        / "external_four_model_1000_postprocess_local"
        / "external_four_model_1000_postprocess"
        / "visual_external4_1000"
        / "crack0001_00_visual_compare.png",
        "图4 外部四模型 1000 样本代表可视化（用于补强外部 baseline 稳定性）",
        width=6.6,
    )


def add_metric_section(doc: Document) -> None:
    """第 2 部分：质量指标对比。"""
    add_heading(doc, "2. 各种算法处理后图像质量指标的比较", 1)
    add_paragraph(
        doc,
        "指标比较采用裂缝区域 EPE、全局 EPE、Dice、边缘保真度和 folding rate 等指标。其中 EPE 和 folding 越低越好，Dice 和 edge fidelity 越高越好。用户原始报告中的 best_epe_v2 作为早期基线，v107 作为当前最新主模型，v108 A-D 作为外部数据消融结果。",
    )

    add_table(
        doc,
        ["模型/实验", "训练轮次", "评估范围", "Crack EPE↓", "Global EPE↓", "Dice↑", "Crack Edge↑", "Global Edge↑", "Folding↓", "结论"],
        [
            ["用户原报告 best_epe_v2", "50 epoch", "120", "110.20", "112.60", "0.0026", "0.3210", "0.3080", "0.5700", "早期基线，已收敛但工程指标明显不足"],
            ["v107 最新主模型", "100 epoch", "120", "15.7639", "16.6204", "0.5278", "0.5534", "0.5517", "0.0222", "当前主结果，几何校正和 folding 显著改善"],
            ["v107 最新主模型", "100 epoch", "1000", "14.8459", "15.0896", "0.5810", "0.5175", "0.5450", "0.0206", "大样本评估仍保持稳定"],
            ["v107 最新主模型", "100 epoch", "全量 10360", "14.6095", "15.2315", "0.5260", "0.3833", "0.4106", "0.0255", "全量结果支撑主实验结论"],
        ],
    )

    add_paragraph(
        doc,
        "与用户原报告 best_epe_v2 相比，v107 在同样 120 样本评估下，Crack EPE 从 110.20px 降至 15.76px，Global EPE 从 112.60px 降至 16.62px，Dice 从 0.0026 提升至 0.5278，folding 从 0.5700 降至 0.0222。说明最新主模型已不再停留在早期百像素级误差，而是在几何校正、裂缝区域对齐和映射稳定性上都有明显提升。",
    )
    add_paragraph(
        doc,
        "外部四模型同时补做了 1000 样本扩展评估，用于验证 baseline 是否只是 120 样本上的偶然波动。结果显示，外部模型在 1000 样本下依然保持稳定，但它们属于 oracle-pair 参考，上界强于我们单图输入模型的设定，不能按完全公平同输入对比来解读。",
    )

    add_table(
        doc,
        ["外部模型", "评估范围", "Crack EPE↓", "Global EPE↓", "Dice↑", "Crack Edge↑", "Folding↓", "说明"],
        [
            ["RAFT", "120", "1.5161", "0.8591", "0.7156", "0.8064", "0.0322", "小样本快速横评结果"],
            ["RAFT", "1000", "1.6220", "0.7900", "0.7308", "0.7900", "0.0303", "外部 baseline 在大样本下仍然稳定"],
            ["GMA/RAFT-small fallback", "120", "2.1497", "1.3111", "0.7082", "0.7479", "0.0328", "小样本快速横评结果"],
            ["GMA/RAFT-small fallback", "1000", "2.3947", "1.2521", "0.7250", "0.7265", "0.0309", "GMA 缺失时的替代参考，结果可用"],
            ["UniMatch", "120", "2.7924", "1.4211", "0.7069", "0.7929", "0.0347", "小样本快速横评结果"],
            ["UniMatch", "1000", "3.8106", "1.6762", "0.7155", "0.7507", "0.0330", "大样本评估后仍保持参考意义"],
            ["SEA-RAFT", "120", "2.8955", "2.0284", "0.7011", "0.7097", "0.0322", "小样本快速横评结果"],
            ["SEA-RAFT", "1000", "3.1400", "1.9537", "0.7188", "0.6942", "0.0304", "同样作为外部上界参考"],
        ],
    )


def add_training_section(doc: Document) -> None:
    """第 3 部分：训练曲线和训练时间。"""
    add_heading(doc, "3. 算法复原水下裂缝图像训练曲线、时间", 1)
    add_paragraph(
        doc,
        "v107 使用最新主模型结构，从 v106 checkpoint 初始化，完整训练 100 epoch。训练总用时约 19 小时 20 分钟。训练末尾 Train Loss 为 0.06034，Train EPE 为 15.973px，Val Loss 为 0.05868，Val EPE 为 15.519px，CrackEPE 为 14.877px。曲线显示 loss、Val EPE 和 Crack EPE 整体持续下降，未见明显过拟合反转。",
    )
    add_image(
        doc,
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "v107_training_curves.png",
        "图5 v107 最新主模型 100 epoch 训练曲线（Loss、EPE、Crack EPE、Learning Rate）",
        width=6.6,
    )
    add_image(
        doc,
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "long_training_curves.png",
        "图6 早期长训练曲线补充（v4-v9）",
        width=6.6,
    )
    add_table(
        doc,
        ["实验", "训练轮次", "训练时间", "最终/关键结果", "说明"],
        [
            ["用户原报告 best_epe", "50 epoch", "约 17.99 小时", "最佳 Val EPE 约 117.885px", "早期基线，训练稳定但指标未达主结果水平"],
            ["v87 长训练", "80 epoch", "长训版本", "Crack EPE 约 24.75px，Dice 约 0.459", "证明长训对几何指标有效，但视觉伪影仍存在"],
            ["v107 最新主模型", "100 epoch", "约 19 小时 20 分钟", "Val EPE 15.519px，CrackEPE 14.877px", "当前主实验结果"],
            ["v108 A-D", "1 epoch smoke", "短周期验证", "用于数据消融方向判断", "不是最终主模型"],
        ],
    )


def add_ablation_section(doc: Document) -> None:
    """第 4 部分：模块消融和外部数据消融。"""
    add_heading(doc, "4. 消融实验，模块增减消融、外部数据消融结果", 1)
    add_paragraph(
        doc,
        "消融实验分为两类：一类是模型/损失模块消融，包括 geometry、anti-ripple、ROI visible-structure、image-space detail head 等；另一类是外部数据消融，用于验证复杂背景、真实水下扰动和不同裂缝类型数据是否能提升结果。",
    )
    add_image(
        doc,
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "late_visual_ablation_metrics_bar.png",
        "图7 后期可视化消融指标对比",
        width=6.6,
    )
    add_image(
        doc,
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "roi_sharpness_win_bar.png",
        "图8 关键 ROI 局部锐度胜出统计",
        width=6.2,
    )

    add_table(
        doc,
        ["消融组", "数据/模块设置", "Crack EPE↓", "Global EPE↓", "Dice↑", "Crack Edge↑", "Global Edge↑", "Folding↓", "结论"],
        [
            ["v108-A", "原训练集 + Roboflow underwater crack", "12.9118", "13.4317", "0.5214", "0.0328", "0.2077", "0.0156", "EPE/folding 最好，但边缘保真明显下降"],
            ["v108-B", "原训练集 + Roboflow concrete / blue crack", "18.6842", "19.4804", "0.4752", "0.4249", "0.4474", "0.0198", "裂缝清楚但水下属性偏弱，不适合作为主数据"],
            ["v108-C", "原训练集 + UIEB/EUVP 水下风格增强", "16.0698", "16.7491", "0.5537", "0.3325", "0.4379", "0.0240", "更适合水下风格增强，不替代裂缝结构数据"],
            ["v108-D", "A+B+C 全部混合", "15.3617", "15.8224", "0.4701", "0.1204", "0.2413", "0.0186", "简单混合不是最优，需要筛选和配比"],
        ],
    )

    add_paragraph(
        doc,
        "模块消融结论：几何主目标和 anti-ripple 约束对稳定复原是必要的；ROI 结构恢复目标可以改善局部结构，但与几何目标存在一定折中；image-space detail head 能解耦局部细节恢复和主 flow，但当前对肉眼可见裂缝锐化的提升仍有限。外部数据消融结论：Roboflow underwater crack 对几何误差有潜力，但直接加入会损伤边缘保真；UIEB/EUVP 更适合作为水下风格增强；concrete/blue crack 与研究对象存在偏差，不建议作为主训练数据。",
    )


def add_conclusion(doc: Document) -> None:
    """添加实验部分总结。"""
    add_heading(doc, "5. 实验结论", 1)
    add_paragraph(
        doc,
        "综合可视化、指标、训练曲线和消融实验，v107 是当前最完整、最可靠的主模型结果。相较用户原始报告中的 best_epe_v2，v107 在裂缝区域 EPE、全局 EPE、Dice、边缘保真和 folding 上均有明显提升，并且已经完成 100 epoch 长训练、1000 样本评估和全量 10360 样本评估。",
    )
    add_paragraph(
        doc,
        "需要客观说明的是，当前实验结果已经能支撑论文实验部分的指标和消融分析，但最终可视化仍不建议表述为完全理想。后续如果继续优化，应以高质量 Roboflow underwater crack 小比例筛选增强、局部 ROI 结构恢复和边缘细节约束为主，而不是简单扩大数据或继续盲目增加训练轮次。",
    )


def main() -> None:
    doc = setup_document()
    title = doc.add_heading("水下裂缝图像复原算法实验部分整理", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(
        doc,
        "本文档根据客户要求整理实验部分，覆盖可视化结果对比、图像质量指标比较、训练曲线与训练时间、模块和外部数据消融实验。主要结果以 v107 最新主模型为主，v108 作为外部数据消融补充。",
    )

    add_visual_section(doc)
    add_metric_section(doc)
    add_training_section(doc)
    add_ablation_section(doc)
    add_conclusion(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
