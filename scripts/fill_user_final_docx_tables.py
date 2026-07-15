#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
只向用户给定的 `实验部分整理最终.docx` 现有表格中填入数据。

注意：
- 不新增表格；
- 不删除表格；
- 不改变表格行列结构；
- 只替换已有单元格中的文字；
- 没有真实实验数据的位置明确写“未单独评估”，避免伪造结果。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


DOCX_PATH = Path(
    r"C:/Users/billy.zhou/Documents/xwechat_files/wxid_bkwta3uyaj7f11_44fc/temp/RWTemp/2026-07/bcb15cc8ba39d84c08631605f8a62f2e/实验部分整理最终.docx"
)


def set_cell_text(cell, text: str, font_size: float = 8.5) -> None:
    """
    在不改变表格结构的前提下替换单元格文字。

    python-docx 的 `cell.text = ...` 会重建单元格内段落。这里采用清空段落 run
    后重新写入的方式，尽量保留单元格、边框、宽度等表格属性。
    """
    paragraph = cell.paragraphs[0]
    for run in list(paragraph.runs):
        run.text = ""
    run = paragraph.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(font_size)

    # 如果单元格原本存在多个段落，仅清空后续段落文字，不删除段落节点，避免影响 Word 结构。
    for extra_para in cell.paragraphs[1:]:
        for extra_run in extra_para.runs:
            extra_run.text = ""


def fill_subjective_table(table) -> None:
    """填写表 0：主观评价表，保留原 6 行 4 列结构。"""
    values = {
        (1, 1): "未单独评估",
        (1, 2): "未单独评估",
        (1, 3): "外部参考/上界；几何校正强，1000样本指标稳定",
        (2, 1): "未单独评估",
        (2, 2): "未单独评估",
        (2, 3): "外部参考/上界；整体稳定，裂缝区域 EPE 略高于 RAFT",
        (3, 1): "未单独评估",
        (3, 2): "未单独评估",
        (3, 3): "外部参考/上界；整体可作为主流 baseline 参考",
        (4, 0): "GMA/RAFT-small fallback",
        (4, 1): "未单独评估",
        (4, 2): "未单独评估",
        (4, 3): "GMA checkpoint 缺失时采用 RAFT-small fallback，结果可用",
        (5, 1): "v108-A 数据消融：EPE/folding 有改善，但边缘保真下降",
        (5, 2): "v108-B 数据消融：裂缝清楚但水下属性偏弱",
        (5, 3): "本文主模型；100 epoch 长训练；几何校正稳定",
    }
    for (r, c), text in values.items():
        set_cell_text(table.cell(r, c), text)


def fill_metric_table(table) -> None:
    """填写表 1：客观指标表，按现有每个数据集一行的结构填入多行指标。"""
    methods = "RAFT\nUniMatch\nSEA-RAFT\nGMA/RAFT-small fallback\nV107算法（our network）"
    for row_idx in [1, 2, 3]:
        set_cell_text(table.cell(row_idx, 1), methods)

    # Roboflow underwater crack：外部四模型没有单独评估；V107 使用 v108-A 数据消融结果。
    set_cell_text(table.cell(1, 2), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n12.9118")
    set_cell_text(table.cell(1, 3), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n13.4317")
    set_cell_text(table.cell(1, 4), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n0.5214")
    set_cell_text(table.cell(1, 5), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n0.0328")
    set_cell_text(table.cell(1, 6), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n0.2077")
    set_cell_text(table.cell(1, 7), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n0.0156")

    # Roboflow concrete / blue crack：外部四模型没有单独评估；V107 使用 v108-B 数据消融结果。
    set_cell_text(table.cell(2, 2), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n18.6842")
    set_cell_text(table.cell(2, 3), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n19.4804")
    set_cell_text(table.cell(2, 4), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n0.4752")
    set_cell_text(table.cell(2, 5), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n0.4249")
    set_cell_text(table.cell(2, 6), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n0.4474")
    set_cell_text(table.cell(2, 7), "未单独评估\n未单独评估\n未单独评估\n未单独评估\n0.0198")

    # 本文数据集：外部四模型 1000 样本 + v107 1000 样本。
    set_cell_text(table.cell(3, 2), "1.6220\n3.8106\n3.1400\n2.3947\n14.8459")
    set_cell_text(table.cell(3, 3), "0.7900\n1.6762\n1.9537\n1.2521\n15.0896")
    set_cell_text(table.cell(3, 4), "0.7308\n0.7155\n0.7188\n0.7250\n0.5810")
    set_cell_text(table.cell(3, 5), "0.7900\n0.7507\n0.6942\n0.7265\n0.5175")
    set_cell_text(table.cell(3, 6), "0.8279\n0.8238\n0.7048\n0.7553\n0.5450")
    set_cell_text(table.cell(3, 7), "0.0303\n0.0330\n0.0304\n0.0309\n0.0206")


def fill_time_table(table) -> None:
    """填写表 2：处理时间/训练设置表，保留原 4 行 6 列结构。"""
    # 两个 Roboflow 数据集没有外部四模型单独评估，因此不填具体耗时。
    for row_idx in [1, 2]:
        for col_idx in [1, 2, 3, 4]:
            set_cell_text(table.cell(row_idx, col_idx), "未单独评估")

    set_cell_text(table.cell(1, 5), "v108-A 1 epoch 数据消融")
    set_cell_text(table.cell(2, 5), "v108-B 1 epoch 数据消融")

    # 本文主数据集：外部模型为推理评估，v107 为本文主模型训练和评估。
    set_cell_text(table.cell(3, 1), "1000样本推理评估完成")
    set_cell_text(table.cell(3, 2), "1000样本推理评估完成")
    set_cell_text(table.cell(3, 3), "1000样本推理评估完成")
    set_cell_text(table.cell(3, 4), "1000样本推理评估完成")
    set_cell_text(table.cell(3, 5), "100 epoch；约19小时20分钟")


def main() -> None:
    """入口：读取用户 Word，填入数据，并直接保存回原文件。"""
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    doc = Document(DOCX_PATH)
    if len(doc.tables) < 3:
        raise RuntimeError(f"模板表格数量不足，当前只有 {len(doc.tables)} 个表格")

    before_shapes = [(len(table.rows), len(table.columns)) for table in doc.tables]

    fill_subjective_table(doc.tables[0])
    fill_metric_table(doc.tables[1])
    fill_time_table(doc.tables[2])

    after_shapes = [(len(table.rows), len(table.columns)) for table in doc.tables]
    if before_shapes != after_shapes:
        raise RuntimeError(f"表格结构发生变化: before={before_shapes}, after={after_shapes}")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)
    print("table_shapes", after_shapes)


if __name__ == "__main__":
    main()
