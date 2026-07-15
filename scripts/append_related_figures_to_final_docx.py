#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
在当前最终 Word 后追加相关图片。

处理原则：
- 不修改已有正文、表格和原有 2 张图片；
- 只在文档末尾追加“补充图”小节；
- 图片采用项目中已经生成并用于客户材料的真实图片；
- 保存前后校验已有表格结构不变，避免破坏用户当前版本。
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "实验部分整理最终_填表补图_补齐指标.docx"


@dataclass(frozen=True)
class FigureSpec:
    """一张需要追加到 Word 的补充图片。"""

    title: str
    path: Path
    width_inches: float = 6.2


FIGURES = [
    FigureSpec(
        "图3 用户原始结果与 v107 主模型指标对比",
        ROOT / "comparison_user_original_vs_v107_assets" / "metric_comparison.png",
    ),
    FigureSpec(
        "图4 后期模块消融指标对比",
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "late_visual_ablation_metrics_bar.png",
    ),
    FigureSpec(
        "图5 v105 / v106 detail-head 诊断曲线",
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "v105_v106_detail_diagnostics.png",
    ),
    FigureSpec(
        "图6 ROI 清晰度 win count 对比",
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "roi_sharpness_win_bar.png",
    ),
    FigureSpec(
        "图7 早期长训练曲线",
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "long_training_curves.png",
    ),
    FigureSpec(
        "图8 Roboflow underwater crack 外部数据代表图",
        ROOT / "delivery_v107_v108_html" / "assets" / "datasets" / "Roboflow_underwater_crack_clear.jpg",
        5.8,
    ),
    FigureSpec(
        "图9 Roboflow concrete / blue crack 外部数据代表图",
        ROOT / "delivery_v107_v108_html" / "assets" / "datasets" / "Roboflow_concrete_crack_clear.jpg",
        5.8,
    ),
]


def add_centered_picture(doc: Document, fig: FigureSpec) -> None:
    """追加一张居中图片和图注。"""

    if not fig.path.exists():
        raise FileNotFoundError(fig.path)

    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(fig.path), width=Inches(fig.width_inches))

    caption = doc.add_paragraph(fig.title)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(9)


def main() -> None:
    """读取当前 Word 并追加相关补充图。"""

    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    doc = Document(DOCX_PATH)
    before_tables = [(len(table.rows), len(table.columns)) for table in doc.tables]
    before_shapes = len(doc.inline_shapes)

    # 避免重复执行脚本时重复追加同一批图片。
    existing_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    if "补充图：消融实验和外部数据相关图片" in existing_text:
        print("related figures already appended; skip")
        return

    doc.add_paragraph("")
    heading = doc.add_paragraph("补充图：消融实验和外部数据相关图片")
    for run in heading.runs:
        run.bold = True
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(12)

    doc.add_paragraph(
        "以下补充图用于支撑消融实验、模块增减消融和外部数据消融结果，"
        "不改变前文已有指标表、训练曲线和可视化对比内容。"
    )

    for fig in FIGURES:
        add_centered_picture(doc, fig)

    doc.save(DOCX_PATH)

    check = Document(DOCX_PATH)
    after_tables = [(len(table.rows), len(table.columns)) for table in check.tables]
    if before_tables != after_tables:
        raise RuntimeError(f"表格结构发生变化: {before_tables} -> {after_tables}")

    print(DOCX_PATH)
    print("tables", after_tables)
    print("inline_shapes_before", before_shapes)
    print("inline_shapes_after", len(check.inline_shapes))


if __name__ == "__main__":
    main()
