#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
在当前最终 Word 中补充消融实验内容。

处理原则：
- 读取项目根目录现有 `实验部分整理最终_填表补图_补齐指标.docx`；
- 不修改已有 3 张对比表和 2 张图片；
- 在文档末尾补充“模块增减消融”和“外部数据消融结果”两张表；
- 写入真实已完成指标，不为未完成实验伪造结果；
- 保存回同一文件，便于用户继续发送当前版本。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "实验部分整理最终_填表补图_补齐指标.docx"


MODULE_ABLATION_ROWS = [
    [
        "v95",
        "High-edge grad 结构强化",
        "120",
        "18.5378",
        "18.8124",
        "0.4990",
        "0.5409",
        "0.5485",
        "0.0165",
        "后期结构强化基线，边缘保真相对较稳",
    ],
    [
        "v100",
        "Anti-ripple / flow 高频约束",
        "120",
        "18.6239",
        "18.7230",
        "0.4977",
        "0.5389",
        "0.5480",
        "0.0183",
        "用于抑制水纹状伪影，几何指标变化不大",
    ],
    [
        "v102",
        "ROI visible-structure + anti-ripple",
        "120",
        "18.5957",
        "18.7035",
        "0.4973",
        "0.5370",
        "0.5478",
        "0.0193",
        "局部结构目标有效但未带来稳定整体提升",
    ],
    [
        "v105",
        "Image-space detail head",
        "120",
        "18.7798",
        "18.8339",
        "0.4950",
        "0.5377",
        "0.5477",
        "0.0161",
        "细节分支与主 flow 解耦，folding 控制较稳",
    ],
    [
        "v106",
        "Strong detail head",
        "120",
        "18.9560",
        "19.1285",
        "0.4931",
        "0.5367",
        "0.5475",
        "0.0153",
        "继续加强细节后，几何误差未继续下降",
    ],
    [
        "v107",
        "最新方法 100 epoch 长训练",
        "120",
        "15.7590",
        "16.6144",
        "0.5278",
        "0.5534",
        "0.5517",
        "0.0222",
        "当前主模型，综合指标最好，作为最终主结果",
    ],
]


DATA_ABLATION_ROWS = [
    [
        "v107 baseline",
        "原训练集，100 epoch 长训练",
        "120",
        "15.7639",
        "16.6204",
        "0.5278",
        "0.5534",
        "0.5517",
        "0.0222",
        "当前主模型基准",
    ],
    [
        "v108-A",
        "原训练集 + Roboflow underwater crack",
        "120",
        "12.9118",
        "13.4317",
        "0.5214",
        "0.0328",
        "0.2077",
        "0.0156",
        "EPE / folding 改善，但边缘保真明显下降",
    ],
    [
        "v108-B",
        "原训练集 + Roboflow concrete / blue crack",
        "120",
        "18.6842",
        "19.4804",
        "0.4752",
        "0.4249",
        "0.4474",
        "0.0198",
        "裂缝形态有补充，但整体指标偏弱",
    ],
    [
        "v108-C",
        "原训练集 + UIEB / EUVP 水下风格增强",
        "120",
        "16.0698",
        "16.7491",
        "0.5537",
        "0.3325",
        "0.4379",
        "0.0240",
        "Dice 最好，但 EPE 和边缘保真未全面优于 v107",
    ],
    [
        "v108-D",
        "A + B + C 混合外部数据",
        "120",
        "15.3617",
        "15.8224",
        "0.4701",
        "0.1204",
        "0.2413",
        "0.0186",
        "EPE / folding 有改善，但 Dice 和边缘保真下降",
    ],
]


def set_font_size(table, size: float = 7.5) -> None:
    """统一表格字体大小，避免列数较多时文字挤出页面。"""

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Microsoft YaHei"
                    run.font.size = Pt(size)


def add_table(doc: Document, title: str, headers: list[str], rows: list[list[str]]) -> None:
    """追加一张带标题的消融实验表。"""

    doc.add_paragraph(title)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_data in rows:
        row = table.add_row()
        for idx, value in enumerate(row_data):
            row.cells[idx].text = value
    set_font_size(table)


def main() -> None:
    """读取当前 Word，补充消融实验内容并保存。"""

    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    doc = Document(DOCX_PATH)
    original_shapes = [(len(table.rows), len(table.columns)) for table in doc.tables[:3]]

    doc.add_paragraph("")
    doc.add_paragraph("消融实验补充说明")
    doc.add_paragraph(
        "消融实验分为两类：第一类是模型模块增减消融，主要观察局部结构约束、"
        "anti-ripple 约束和 detail head 对几何误差、边缘保真和 folding 的影响；"
        "第二类是外部数据消融，主要观察 Roboflow、UIEB/EUVP 等新增数据对主模型的影响。"
        "其中 v107 为当前最终主模型，v108 A-D 为短周期数据消融实验，不作为最终主模型替代。"
    )

    headers = [
        "版本",
        "模块/设置",
        "样本数",
        "Crack EPE↓",
        "Global EPE↓",
        "Dice↑",
        "Crack Edge↑",
        "Global Edge↑",
        "Folding↓",
        "结论",
    ]
    add_table(doc, "表3 模块增减消融实验结果", headers, MODULE_ABLATION_ROWS)

    doc.add_paragraph(
        "从模块消融结果看，v95 到 v106 的后期模块主要用于抑制局部伪影、增强 ROI 结构和测试细节分支。"
        "这些模块对 folding 和局部稳定性有一定帮助，但单独增强局部细节并没有稳定降低 Crack EPE。"
        "最终 v107 通过最新方法进行 100 epoch 长训练后，Crack EPE、Global EPE、Dice 和边缘保真均明显优于后期短周期消融版本，"
        "因此最终主结果仍采用 v107。"
    )

    add_table(doc, "表4 外部数据消融实验结果", headers, DATA_ABLATION_ROWS)

    doc.add_paragraph(
        "从外部数据消融结果看，Roboflow underwater crack 和混合外部数据可以降低部分几何误差和 folding，"
        "但会明显牺牲边缘保真；UIEB/EUVP 风格增强能提高 Dice，但不能同时改善 EPE、边缘保真和 folding。"
        "因此新增外部数据目前可以作为消融分析和数据质量讨论的依据，但不建议直接替代原训练集上的 v107 主模型。"
    )

    after_shapes = [(len(table.rows), len(table.columns)) for table in doc.tables[:3]]
    if original_shapes != after_shapes:
        raise RuntimeError(f"已有前三张表结构发生变化: {original_shapes} -> {after_shapes}")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)
    print("original_tables", original_shapes)
    print("all_tables", [(len(table.rows), len(table.columns)) for table in Document(DOCX_PATH).tables])
    print("inline_shapes", len(Document(DOCX_PATH).inline_shapes))


if __name__ == "__main__":
    main()
