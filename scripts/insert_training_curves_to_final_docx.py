#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
在最终 Word 的“训练曲线”章节内补充多组训练/消融曲线。

背景：
用户指出当前“主训练数据集上，各种方法的训练曲线”部分看起来只放了 v107。
本脚本不改动已有内容，只在“图2 各种模型训练曲线”后追加说明和相关曲线：
1. v107 最新方法 100 epoch 长训练曲线；
2. v4-v9 早期长训练曲线；
3. v95-v106 后期模块消融指标曲线；
4. v105/v106 detail-head 诊断曲线。

说明：
- RAFT、UniMatch、SEA-RAFT、GMA/RAFT-small fallback 是外部预训练推理 baseline，
  本项目没有重新训练这些模型，因此没有本项目训练曲线；
- v108 A-D 属于 1 epoch 数据消融 smoke，不是长训练曲线；
- 脚本会检查标记文本，避免重复插入。
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "实验部分整理最终_填表补图_补齐指标.docx"


@dataclass(frozen=True)
class CurveFigure:
    """训练曲线补充图配置。"""

    caption: str
    path: Path
    width_inches: float = 6.2


CURVE_FIGURES = [
    CurveFigure(
        "图2-1 v107 最新方法 100 epoch 完整训练曲线",
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "v107_training_curves.png",
    ),
    CurveFigure(
        "图2-2 v4-v9 早期长训练曲线",
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "long_training_curves.png",
    ),
    CurveFigure(
        "图2-3 v95-v106 后期模块消融指标曲线",
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "late_visual_ablation_metrics_bar.png",
    ),
    CurveFigure(
        "图2-4 v105/v106 detail-head 诊断曲线",
        ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "v105_v106_detail_diagnostics.png",
    ),
]


def style_paragraph(paragraph, font_size: float = 10.5, bold: bool = False) -> None:
    """统一设置新增段落字体，避免和原文风格差异太大。"""

    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(font_size)
        run.bold = bold


def insert_after(anchor, new_paragraph):
    """把新段落移动到 anchor 段落之后，并返回新段落作为新的锚点。"""

    anchor._p.addnext(new_paragraph._p)  # pylint: disable=protected-access
    return new_paragraph


def add_paragraph_after(doc: Document, anchor, text: str, font_size: float = 10.5, bold: bool = False):
    """在指定段落后插入普通文字段落。"""

    paragraph = doc.add_paragraph(text)
    style_paragraph(paragraph, font_size=font_size, bold=bold)
    return insert_after(anchor, paragraph)


def add_picture_after(doc: Document, anchor, fig: CurveFigure):
    """在指定段落后插入居中图片和图注。"""

    if not fig.path.exists():
        raise FileNotFoundError(fig.path)

    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(str(fig.path), width=Inches(fig.width_inches))
    anchor = insert_after(anchor, pic_para)

    caption = doc.add_paragraph(fig.caption)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(caption, font_size=9)
    return insert_after(anchor, caption)


def find_training_anchor(doc: Document):
    """定位“图2 各种模型训练曲线”段落，作为补充内容插入点。"""

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("图2") and "训练曲线" in text:
            return paragraph
    raise RuntimeError("未找到“图2 各种模型训练曲线”段落")


def main() -> None:
    """主入口：在训练曲线章节插入补充曲线。"""

    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    doc = Document(DOCX_PATH)
    before_tables = [(len(table.rows), len(table.columns)) for table in doc.tables]
    before_shapes = len(doc.inline_shapes)

    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    marker = "训练曲线补充说明"
    if marker in all_text:
        print("training curve supplement already exists; skip")
        return

    anchor = find_training_anchor(doc)
    anchor = add_paragraph_after(doc, anchor, "", font_size=10.5)
    anchor = add_paragraph_after(doc, anchor, marker, font_size=11, bold=True)
    anchor = add_paragraph_after(
        doc,
        anchor,
        "外部 RAFT、UniMatch、SEA-RAFT、GMA/RAFT-small fallback 使用公开预训练权重做推理对比，"
        "本项目没有重新训练这些外部模型，因此不提供本项目训练曲线。"
        "本文训练曲线主要覆盖内部模型的长训练和后期消融过程：v107 为最新方法 100 epoch 长训练，"
        "v4-v9 为早期长训练主线，v95-v106 为后期模块/可视化消融，v105/v106 为 detail-head 诊断。"
        "v108 A-D 是 1 epoch 外部数据消融 smoke，只用于数据消融方向判断，不作为长训练曲线展示。",
        font_size=10,
    )

    for fig in CURVE_FIGURES:
        anchor = add_picture_after(doc, anchor, fig)

    doc.save(DOCX_PATH)

    check = Document(DOCX_PATH)
    after_tables = [(len(table.rows), len(table.columns)) for table in check.tables]
    if before_tables != after_tables:
        raise RuntimeError(f"表格结构发生变化: {before_tables} -> {after_tables}")

    print(DOCX_PATH)
    print("tables", after_tables)
    print("inline_shapes_before", before_shapes)
    print("inline_shapes_after", len(check.inline_shapes))


if __name__ == "__main__":
    main()
