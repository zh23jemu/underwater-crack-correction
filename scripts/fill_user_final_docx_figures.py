#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
向用户给定的 `实验部分整理最终.docx` 补充图1和图2。

处理原则：
- 不修改已有表格结构；
- 不重写整篇文档；
- 只在已有“图1”“图2”图注附近插入图片；
- 如果已经插入过图片，不重复插入，避免多次运行造成重复图。
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


DOCX_PATH = Path(
    r"C:/Users/billy.zhou/Documents/xwechat_files/wxid_bkwta3uyaj7f11_44fc/temp/RWTemp/2026-07/bcb15cc8ba39d84c08631605f8a62f2e/实验部分整理最终.docx"
)
ROOT = Path(__file__).resolve().parents[1]

FIG1 = (
    ROOT
    / "external_four_model_1000_postprocess_local"
    / "external_four_model_1000_postprocess"
    / "visual_external4_1000"
    / "crack0001_00_visual_compare.png"
)
FIG2 = ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "v107_training_curves.png"


def insert_picture_before(paragraph, image_path: Path, width_inches: float) -> None:
    """在指定段落前插入一张居中图片。"""
    if not image_path.exists():
        raise FileNotFoundError(image_path)

    new_p = paragraph.insert_paragraph_before()
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def main() -> None:
    """读取用户 Word，在图1和图2图注前补图片，并保持表格结构不变。"""
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    doc = Document(DOCX_PATH)
    table_shapes_before = [(len(table.rows), len(table.columns)) for table in doc.tables]

    # 当前用户文件最初没有图片；如果已经补过图，则不重复插入。
    if len(doc.inline_shapes) < 2:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text.startswith("图1"):
                insert_picture_before(paragraph, FIG1, 6.4)
            elif text.startswith("图2"):
                insert_picture_before(paragraph, FIG2, 6.4)

    table_shapes_after = [(len(table.rows), len(table.columns)) for table in doc.tables]
    if table_shapes_before != table_shapes_after:
        raise RuntimeError(f"表格结构发生变化: before={table_shapes_before}, after={table_shapes_after}")

    doc.save(DOCX_PATH)
    print(DOCX_PATH)
    print("table_shapes", table_shapes_after)
    print("inline_shapes", len(Document(DOCX_PATH).inline_shapes))


if __name__ == "__main__":
    main()
