#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
把漏掉的 v87 80 epoch 长训版本补充到最终 Word 的训练曲线章节。

说明：
- 若本地已从远端找回 v87 epoch 汇总行，则生成真实的 80 epoch 训练曲线；
- 若 epoch 汇总行缺失，则退回使用项目记录中已保存的真实 v87 训练末尾、最佳值和
  120 样本评估指标，生成“v87 80 epoch 长训指标摘要图”；
- 插入 Word 时只追加段落和图片，不修改已有表格、已有图片和已有正文。
"""

from __future__ import annotations

from pathlib import Path
import re

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "实验部分整理最终_填表补图_补齐指标.docx"
EPOCH_LOG_PATH = ROOT / "output_crackwarp_slurm" / "v87_spatial_cap_long80_from_v83" / "v87_epoch_summary_lines.log"
FIG_PATH = ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "v87_long80_training_curves.png"
OLD_FIG_PATH = ROOT / "delivery_v107_v108_html" / "assets" / "curves" / "v87_long80_summary.png"
NEW_MARKER = "图2-5 v87 80 epoch 长训练曲线"
OLD_MARKER = "图2-5 v87 80 epoch 长训指标摘要"


def style_paragraph(paragraph, font_size: float = 10.0, bold: bool = False) -> None:
    """统一新增文字样式，使补充内容和现有 Word 风格接近。"""

    for run in paragraph.runs:
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(font_size)
        run.bold = bold


def insert_after(anchor, new_paragraph):
    """把新段落移动到 anchor 段落之后，并返回新段落作为新的锚点。"""

    anchor._p.addnext(new_paragraph._p)  # pylint: disable=protected-access
    return new_paragraph


def add_text_after(doc: Document, anchor, text: str, font_size: float = 10.0, bold: bool = False):
    """在指定段落后追加文本段落。"""

    paragraph = doc.add_paragraph(text)
    style_paragraph(paragraph, font_size=font_size, bold=bold)
    return insert_after(anchor, paragraph)


def add_picture_after(doc: Document, anchor, image_path: Path, caption: str):
    """在指定段落后追加居中图片和图注。"""

    if not image_path.exists():
        raise FileNotFoundError(image_path)

    pic_para = doc.add_paragraph()
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_para.add_run().add_picture(str(image_path), width=Inches(6.2))
    anchor = insert_after(anchor, pic_para)

    caption_para = doc.add_paragraph(caption)
    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_paragraph(caption_para, font_size=9.0)
    return insert_after(anchor, caption_para)


def parse_epoch_summary_lines() -> list[dict[str, float]]:
    """解析从远端 train.log 抽取的 v87 epoch 汇总行。"""

    if not EPOCH_LOG_PATH.exists():
        return []

    pattern = re.compile(
        r"Epoch (?P<epoch>\d+)/80 \| Train Loss: (?P<train_loss>[0-9.]+) "
        r"EPE: (?P<train_epe>[0-9.]+)px \| Val Loss: (?P<val_loss>[0-9.]+) "
        r"EPE: (?P<val_epe>[0-9.]+)px CrackEPE: (?P<crack_epe>[0-9.]+)px "
        r"\| LR: (?P<lr>[0-9.eE+-]+)"
    )
    rows: list[dict[str, float]] = []
    for line in EPOCH_LOG_PATH.read_text(encoding="utf-8", errors="replace").splitlines():
        match = pattern.search(line)
        if not match:
            continue
        item = {key: float(value) for key, value in match.groupdict().items()}
        rows.append(item)
    return rows


def build_v87_training_curve() -> bool:
    """优先用真实 epoch 汇总行生成 v87 80 epoch 训练曲线。"""

    rows = parse_epoch_summary_lines()
    if len(rows) != 80:
        return False

    FIG_PATH.parent.mkdir(parents=True, exist_ok=True)
    epochs = [row["epoch"] for row in rows]

    fig, axes = plt.subplots(2, 2, figsize=(12.8, 7.2), dpi=180)
    fig.patch.set_facecolor("white")

    axes[0, 0].plot(epochs, [row["train_loss"] for row in rows], label="Train Loss", color="#2563eb", linewidth=2)
    axes[0, 0].plot(epochs, [row["val_loss"] for row in rows], label="Val Loss", color="#f97316", linewidth=2)
    axes[0, 0].set_title("Loss over 80 epochs")
    axes[0, 0].set_xlabel("Epoch")
    axes[0, 0].set_ylabel("Loss")
    axes[0, 0].grid(alpha=0.25)
    axes[0, 0].legend()

    axes[0, 1].plot(epochs, [row["train_epe"] for row in rows], label="Train EPE", color="#0f766e", linewidth=2)
    axes[0, 1].plot(epochs, [row["val_epe"] for row in rows], label="Val EPE", color="#dc2626", linewidth=2)
    axes[0, 1].plot(epochs, [row["crack_epe"] for row in rows], label="Crack EPE", color="#7c3aed", linewidth=2)
    axes[0, 1].set_title("EPE metrics over 80 epochs")
    axes[0, 1].set_xlabel("Epoch")
    axes[0, 1].set_ylabel("px")
    axes[0, 1].grid(alpha=0.25)
    axes[0, 1].legend()

    axes[1, 0].plot(epochs, [row["lr"] for row in rows], color="#0891b2", linewidth=2)
    axes[1, 0].set_title("Learning rate schedule")
    axes[1, 0].set_xlabel("Epoch")
    axes[1, 0].set_ylabel("LR")
    axes[1, 0].grid(alpha=0.25)

    final = rows[-1]
    summary_names = ["Train EPE", "Val EPE", "Crack EPE"]
    summary_values = [final["train_epe"], final["val_epe"], final["crack_epe"]]
    axes[1, 1].bar(summary_names, summary_values, color=["#2563eb", "#f97316", "#7c3aed"])
    axes[1, 1].set_title("Final epoch metrics")
    axes[1, 1].set_ylabel("px")
    axes[1, 1].grid(axis="y", alpha=0.25)
    for index, value in enumerate(summary_values):
        axes[1, 1].text(index, value + 0.35, f"{value:.2f}", ha="center", va="bottom", fontsize=9)

    fig.suptitle("v87 80 epoch long training curves", fontsize=14, y=0.99)
    fig.tight_layout(rect=(0, 0, 1, 0.96))
    fig.savefig(FIG_PATH, bbox_inches="tight")
    plt.close(fig)
    return True


def build_v87_summary_figure() -> None:
    """用真实记录指标生成 v87 80 epoch 长训摘要图。"""

    OLD_FIG_PATH.parent.mkdir(parents=True, exist_ok=True)

    # 这些数值来自项目 AGENTS.md 中记录的 v87 长训结果和 120 样本评估结果。
    train_metrics = {
        "Train EPE\nfinal": 23.768,
        "Val EPE\nfinal": 24.028,
        "Best Val\nEPE": 24.015,
        "Best Crack\nEPE": 22.817,
    }
    eval_metrics = {
        "Crack EPE": 24.747450,
        "Global EPE": 24.789375,
        "Dice x50": 0.459496 * 50,
        "Crack Edge x50": 0.527016 * 50,
        "Folding x500": 0.059502 * 500,
    }

    fig, axes = plt.subplots(1, 2, figsize=(13.5, 4.9), dpi=180)
    fig.patch.set_facecolor("white")

    colors_left = ["#3b82f6", "#2563eb", "#0f766e", "#14b8a6"]
    axes[0].bar(list(train_metrics.keys()), list(train_metrics.values()), color=colors_left)
    axes[0].set_title("v87 long training summary (80 epochs)")
    axes[0].set_ylabel("px")
    axes[0].grid(axis="y", alpha=0.25)
    for index, value in enumerate(train_metrics.values()):
        axes[0].text(index, value + 0.6, f"{value:.2f}", ha="center", va="bottom", fontsize=8)

    colors_right = ["#ef4444", "#f97316", "#22c55e", "#84cc16", "#8b5cf6"]
    axes[1].bar(list(eval_metrics.keys()), list(eval_metrics.values()), color=colors_right)
    axes[1].set_title("120-sample evaluation after v87 long training")
    axes[1].set_ylabel("scaled value")
    axes[1].grid(axis="y", alpha=0.25)
    axes[1].tick_params(axis="x", labelrotation=18)
    labels = ["24.75", "24.79", "0.459", "0.527", "0.0595"]
    for index, (value, label) in enumerate(zip(eval_metrics.values(), labels)):
        axes[1].text(index, value + 0.7, label, ha="center", va="bottom", fontsize=8)

    note = (
        "Note: local package contains v87 final/best/eval metrics, but no per-epoch train.log. "
        "This figure is a long-training metric summary, not a fabricated epoch curve."
    )
    fig.text(0.5, 0.02, note, ha="center", va="bottom", fontsize=8, color="#475569")
    fig.tight_layout(rect=(0, 0.07, 1, 1))
    fig.savefig(OLD_FIG_PATH, bbox_inches="tight")
    plt.close(fig)


def find_insert_anchor(doc: Document):
    """优先插入到图2-4后；如果不存在，则退回插入到图2训练曲线后。"""

    fallback = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("图2") and "训练曲线" in text:
            fallback = paragraph
        if text.startswith("图2-4"):
            return paragraph
    if fallback is not None:
        return fallback
    raise RuntimeError("未找到训练曲线插入位置")


def refresh_existing_figure(doc: Document, marker: str, new_caption: str, image_path: Path) -> bool:
    """如果 Word 中已存在图2-5，则替换其前一段图片的嵌入数据。"""

    marker_index = None
    for index, paragraph in enumerate(doc.paragraphs):
        if marker in paragraph.text:
            marker_index = index
            paragraph.text = new_caption
            style_paragraph(paragraph, font_size=9.0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break
    if marker_index is None:
        return False

    image_bytes = image_path.read_bytes()
    for paragraph in reversed(doc.paragraphs[:marker_index]):
        blips = paragraph._p.xpath(".//a:blip")  # pylint: disable=protected-access
        if not blips:
            continue
        embed_id = blips[0].get(qn("r:embed"))
        if not embed_id:
            continue
        image_part = doc.part.related_parts[embed_id]
        image_part._blob = image_bytes  # pylint: disable=protected-access
        return True

    raise RuntimeError("找到图2-5图注，但未找到对应的前置图片")


def main() -> None:
    """生成摘要图并补入 Word，同时校验表格结构不变。"""

    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    has_epoch_curve = build_v87_training_curve()
    if not has_epoch_curve:
        build_v87_summary_figure()

    doc = Document(DOCX_PATH)
    before_tables = [(len(table.rows), len(table.columns)) for table in doc.tables]
    before_shapes = len(doc.inline_shapes)

    marker = NEW_MARKER if has_epoch_curve else OLD_MARKER
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    existing_marker = NEW_MARKER if NEW_MARKER in all_text else OLD_MARKER if OLD_MARKER in all_text else ""
    if existing_marker:
        refreshed = refresh_existing_figure(doc, existing_marker, marker, FIG_PATH if has_epoch_curve else OLD_FIG_PATH)
        doc.save(DOCX_PATH)
        print("v87 long80 supplement already exists; refreshed", refreshed)
        return

    anchor = find_insert_anchor(doc)
    anchor = add_text_after(
        doc,
        anchor,
        "补充说明：v87 是此前已经完成的 80 epoch 长训版本，训练轮次不同于 v108 的 1 epoch 数据消融 smoke。"
        "该版本证明长训能明显改善常规指标，但 ROI 复核仍存在局部涟漪、糊化和形变伪影，因此只作为长训代表和诊断基线。",
        font_size=10.0,
    )
    add_picture_after(doc, anchor, FIG_PATH if has_epoch_curve else OLD_FIG_PATH, marker)

    doc.save(DOCX_PATH)

    check = Document(DOCX_PATH)
    after_tables = [(len(table.rows), len(table.columns)) for table in check.tables]
    if before_tables != after_tables:
        raise RuntimeError(f"表格结构发生变化: {before_tables} -> {after_tables}")

    print(DOCX_PATH)
    print(FIG_PATH if has_epoch_curve else OLD_FIG_PATH)
    print("tables", after_tables)
    print("inline_shapes_before", before_shapes)
    print("inline_shapes_after", len(check.inline_shapes))


if __name__ == "__main__":
    main()
